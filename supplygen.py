import pandas as pd
from docx import Document
from docx.shared import Inches


import pyalex
from pyalex import Works

def get_openalex_search(LLMstr):
    
    keywords=LLMstr.split(',')
    
    # 配置你的电子邮件
    pyalex.config.email = "1017453548@qq.com"

    # 构建查询字符串
    query = " ".join(keywords)

    # 使用pyalex进行搜索
    results = Works().search(query).get()


    # 使用pyalex进行搜索，并按时间排序
    # results = Works().search(query).sort(publication_date="desc").get()

    # 初始化结果列表
    results_list = []

    # 遍历每个结果
    for result in results:
        try:
            title = result['title']
            authors = ', '.join([author['author']['display_name'] for author in result['authorships']])
            abstract = result['abstract']
            doi = result['doi']
            publication_date = result['publication_date']
            journal = result['primary_location']['source']['display_name']

            # 构建信息字典
            result_info = {
                'Title': title,
                'Authors': authors,
                'Abstract': abstract,
                'DOI': doi,
                'Publication Date': publication_date,
                'Journal': journal
            }

            # 添加到结果列表
            results_list.append(result_info)
        except Exception as e:
            print(f"Error processing result: {e}")
            continue

    return results_list




import pandas as pd
import numpy as np




####文件树生成####

# import pandas as pd
# import os
# from pathlib import Path

# def process_reactions(input_file, output_dir="output"):
#     # 读取前1000行数据
#     try:
#         df = pd.read_excel(input_file, nrows=1000)
#     except FileNotFoundError:
#         print(f"错误：文件 {input_file} 不存在")
#         return

#     # 校验必要列存在
#     required_columns = ['Reaction ID', 'Reaction', 'References']
#     if not set(required_columns).issubset(df.columns):
#         missing = set(required_columns) - set(df.columns)
#         print(f"错误：缺少必要列 {missing}")
#         return

#     # 准备输出目录
#     Path(output_dir).mkdir(exist_ok=True)
#     input_contents = []

#     for index, row in df.iterrows():
#         try:
#             # 清洗 Reaction ID
#             reaction_id = str(row['Reaction ID']).strip()
#             reaction_id = "".join(c for c in reaction_id if c not in r'\/:*?"<>|')
            
#             if not reaction_id:
#                 print(f"第 {index+1} 行：Reaction ID 为空，已跳过")
#                 continue

#             # 创建反应目录
#             reaction_dir = Path(output_dir) / reaction_id
#             reaction_dir.mkdir(exist_ok=True)

#             # 写入 Reaction 文件
#             reaction = str(row['Reaction']).strip()
#             with open(reaction_dir/'Reaction.txt', 'w', encoding='utf-8') as f:
#                 f.write(reaction)

#             # 写入 References 文件
#             references = str(row['References']).strip()
#             with open(reaction_dir/'Reference.txt', 'w', encoding='utf-8') as f:
#                 f.write(references)

#             # 收集 input 内容
#             processed = reaction.split('>>')[-1].strip()
#             if processed:
#                 input_contents.append(processed)
#             with open(reaction_dir/'input.txt', 'w', encoding='utf-8') as f:
#                 f.write(processed)

#         except Exception as e:
#             print(f"处理第 {index+1} 行时出错：{str(e)}")
#             continue

#     # 写入全局 input.txt
#     with open(Path(output_dir)/'input.txt', 'w', encoding='utf-8') as f:
#         f.write("\n".join(input_contents))

#     print(f"处理完成，共成功处理 {len(input_contents)} 条有效记录")

# # 使用示例
# if __name__ == "__main__":
#     process_reactions(
#         input_file="validated_output.xlsx",  # 替换为实际文件路径
#         output_dir="reactions_data"
#     )

####通过get请求获取edit的预测反应路径，并组合成反应式（top2）

# import os
# import shutil
# import requests
# import time
# from pathlib import Path

# def process_reactions_api(api_endpoint, 
#                          source_dir="reactions_data",
#                          temp_dir="./api_temp",
#                          max_retries=1):
#     """
#     带自动重试机制的 API 处理流程
#     :param api_endpoint: API 调用命令/URL
#     :param source_dir: Reaction ID 文件夹根目录
#     :param temp_dir: 临时工作目录
#     :param max_retries: 最大重试次数
#     """
#     # 创建临时目录
#     Path(temp_dir).mkdir(parents=True, exist_ok=True)
    
#     processed = 0
#     error_count = 0
#     reaction_dirs = [d for d in os.listdir(source_dir) 
#                     if os.path.isdir(os.path.join(source_dir, d))]
#     print(reaction_dirs)
    
#     for reaction_id in reaction_dirs:
#         current_dir = Path(source_dir) / reaction_id
#         input_file = current_dir / "input.txt"
#         output_file = current_dir / "ranked_preds.txt"
        
#         if not input_file.exists():
#             print(f"跳过 {reaction_id}：缺少 input.txt")
#             continue

#         # 清理临时目录
#         shutil.rmtree(temp_dir, ignore_errors=True)
#         Path(temp_dir).mkdir()

#         try:
#             # 阶段1：文件准备
#             temp_input = Path(temp_dir) / "input.txt"
#             shutil.copy(input_file, temp_input)
            
#             # 阶段2：API 调用
#             for attempt in range(max_retries):
#                 try:
#                     # 如果是 HTTP API
#                     response = requests.get(
#                         url=api_endpoint,
#                         files={'file': open(temp_input, 'rb')},
#                         timeout=300  # 5 分钟超时
#                     )
#                     response.raise_for_status()
                    
#                     # 保存 API 响应
#                     with open(Path(temp_dir)/"api.txt", "wb") as f:
#                         f.write(response.content)
#                     break
#                 except Exception as e:
#                     if attempt == max_retries -1:
#                         raise
#                     print(f"{reaction_id} 第 {attempt+1} 次重试...")
#                     time.sleep(2 ** attempt)  # 指数退避

#             # 阶段3：结果处理
#             result_file = Path(temp_dir)/"ranked_preds.txt"
#             if result_file.exists() and result_file.stat().st_size > 0:
#                 shutil.copy(result_file, output_file)
#                 print(f"√ 成功处理 {reaction_id}")
#                 processed +=1
#             else:
#                 raise Exception("API 未返回有效结果文件")

#         except Exception as e:
#             print(f"× 处理 {reaction_id} 失败：{str(e)}")
#             error_count +=1

#     print(f"\n处理完成：成功 {processed} 条，失败 {error_count} 条")

# # 使用示例（HTTP API 版本）
# if __name__ == "__main__":
#     process_reactions_api(
#         api_endpoint="http://localhost:8005/run-script",  # 替换实际 API 地址
#         source_dir="reactions_data",
#         temp_dir="E:/dockerdata/api",
#         max_retries=1
#     )

####通过rdkit绘制反应路径

from rdkit import Chem
from rdkit.Chem import AllChem, Draw
from rdkit.Chem.Draw import rdMolDraw2D
import os

def draw_reaction_scheme(
    reaction_smarts: str, 
    output_file: str = "reaction.png",
    image_size: tuple = (600, 250)
) -> dict:
    """
    根据SMARTS反应式生成化学反应示意图
    
    参数:
        reaction_smarts (str): SMARTS格式的反应式字符串
        output_file (str): 输出图片路径，默认当前目录的reaction.png
        image_size (tuple): 图片尺寸，默认(600,250)
    
    返回:
        dict: 包含执行状态和文件路径的字典
    """
    try:
        # 验证输入参数
        if not reaction_smarts:
            raise ValueError("反应式不能为空")
            
        if not output_file.endswith('.png'):
            output_file += '.png'

        # 创建反应对象
        rxn = AllChem.ReactionFromSmarts(reaction_smarts, useSmiles=True)
        
        # 创建绘图对象
        drawer = Draw.MolDraw2DCairo(*image_size)
        
        # 绘制反应式
        drawer.DrawReaction(rxn)
        drawer.FinishDrawing()

        # 保存图片
        png_data = drawer.GetDrawingText()
        with open(output_file, 'wb') as f:
            f.write(png_data)

        return {
            "status": "success",
            "message": "反应式示意图生成成功",
            "output_path": os.path.abspath(output_file)
        }

    except (Chem.AtomValenceException, Chem.KekulizeException) as e:
        return {
            "status": "error",
            "message": f"化学结构解析失败: {str(e)}",
            "output_path": None
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"未知错误: {str(e)}",
            "output_path": None
        }
        
source_dir = "reactions_data"
reaction_dirs = [d for d in os.listdir(source_dir) 
                if os.path.isdir(os.path.join(source_dir, d))]
print(reaction_dirs)

# for reaction_id in reaction_dirs:
#     current_dir = os.path.join(source_dir, reaction_id)
    
#     # 检查必需文件
#     ranked_preds = os.path.join(current_dir, "ranked_preds.txt")
#     input_file = os.path.join(current_dir, "input.txt")
    
#     if not os.path.exists(ranked_preds):
#         print(f"跳过 {reaction_id}：缺少 ranked_preds.txt")
#         continue

#     try:
#         # 读取input内容
#         with open(input_file, 'r') as f:
#             input_smiles = f.read().strip()

#         # 生成基础反应式图 r.png
#         base_reaction = f"{input_smiles}>>{input_smiles}"
#         draw_reaction_scheme(
#             base_reaction,
#             output_file=os.path.join(current_dir, "r.png"),
#             image_size=(800, 400)
#         )

#         # 处理预测结果
#         with open(ranked_preds, 'r') as f:
#             pred_lines = [line.strip() for line in f.readlines()[:5]]  # 取前5行

#         for idx, pred_smiles in enumerate(pred_lines, 1):
#             if not pred_smiles:
#                 continue

#             # 构建完整反应式
#             full_reaction = f"{pred_smiles}>>{input_smiles}"
            
#             # 生成编号图片
#             result = draw_reaction_scheme(
#                 full_reaction,
#                 output_file=os.path.join(current_dir, f"{idx}.png"),
#                 image_size=(1200, 400)
#             )

#             # 处理绘图结果
#             if result["status"] == "error":
#                 print(f"{reaction_id} 第{idx}个预测式绘图失败: {result['message']}")

#     except Exception as e:
#         print(f"处理 {reaction_id} 时发生异常: {str(e)}")
#         continue

######再处理##########
# for reaction_id in reaction_dirs:
#     current_dir = os.path.join(source_dir, reaction_id)
    
#     # 检查必需文件
#     required_files = {
#         "reaction": os.path.join(current_dir, "reaction.txt"),
#         "ranked_preds": os.path.join(current_dir, "ranked_preds.txt"),
#         "input": os.path.join(current_dir, "input.txt")
#     }
    
#     # 验证文件存在性
#     missing_files = [k for k, v in required_files.items() if not os.path.exists(v)]
#     if missing_files:
#         print(f"跳过 {reaction_id}：缺少 {', '.join(missing_files)}")
#         continue

#     try:
#         # 读取原始反应式
#         with open(required_files["reaction"], 'r') as f:
#             original_reaction = f.read().strip()
        
#         # 生成原始反应式示意图
#         result = draw_reaction_scheme(
#             original_reaction,
#             output_file=os.path.join(current_dir, "r2.png"),
#             image_size=(1000, 400)
#         )
#         if result["status"] == "error":
#             print(f"{reaction_id} 原始反应式绘图失败: {result['message']}")

#         # # 读取输入内容和预测结果
#         # with open(required_files["input"], 'r') as f:
#         #     target_smiles = f.read().strip()
            
#         # with open(required_files["ranked_preds"], 'r') as f:
#         #     pred_lines = [line.strip() for line in f.readlines()[:5]]

#         # # 生成预测反应式示意图
#         # for idx, precursor in enumerate(pred_lines, 1):
#         #     if not precursor or '>' in precursor:  # 过滤无效行
#         #         continue
                
#         #     full_reaction = f"{precursor}>>{target_smiles}"
            
#         #     result = draw_reaction_scheme(
#         #         full_reaction,
#         #         output_file=os.path.join(current_dir, f"{idx}.png"),
#         #         image_size=(1200, 400)
#         #     )
            
#         #     if result["status"] == "error":
#         #         print(f"{reaction_id} 预测反应式{idx}绘图失败: {result['message']}")

#     except Exception as e:
#         print(f"处理 {reaction_id} 时发生异常: {str(e)}")
#         continue
    
####将所有信息插入到word文档中##
from rdkit import Chem
from rdkit.Chem import AllChem

def normalize_smiles(smiles):
    """标准化SMILES字符串"""
    mol = Chem.MolFromSmiles(smiles)
    if mol:
        return Chem.MolToSmiles(mol, isomericSmiles=True, canonical=True)
    return None

def are_reactions_equal(rxn1_smarts, rxn2_smarts):
    """
    判断两个反应式是否化学等价
    
    参数：
    rxn1_smarts (str): 第一个反应式（SMARTS格式）
    rxn2_smarts (str): 第二个反应式（SMARTS格式）
    
    返回：
    bool: True表示化学等价，False表示不等价
    """
    try:
        # 解析反应式
        rxn1 = AllChem.ReactionFromSmarts(rxn1_smarts)
        rxn2 = AllChem.ReactionFromSmarts(rxn2_smarts)
        
        # 比较反应式组成
        def get_components(rxn):
            return (
                sorted([normalize_smiles(Chem.MolToSmiles(m)) for m in rxn.GetReactants()]),
                sorted([normalize_smiles(Chem.MolToSmiles(m)) for m in rxn.GetProducts()]),
                sorted([normalize_smiles(Chem.MolToSmiles(m)) for m in rxn.GetAgents()])
            )
        
        r1_react, r1_prod, r1_agent = get_components(rxn1)
        r2_react, r2_prod, r2_agent = get_components(rxn2)
        
        return (
            r1_react == r2_react and
            r1_prod == r2_prod and
            r1_agent == r2_agent
        )
        
    except Exception as e:
        print(f"反应式解析失败: {str(e)}")
        return False
    

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def generate_report(source_dir="reactions_data", output_file="report.docx"):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    index=0
    matchcount=0
    nomatchcount=0
    for reaction_id in os.listdir(source_dir):
        
        current_dir = os.path.join(source_dir, reaction_id)
        if not os.path.isdir(current_dir):
            continue

        # 检查必需文件
        required_files = {
            "reference": os.path.join(current_dir, "Reference.txt"),
            "reaction_img": os.path.join(current_dir, "r2.png"),
            "preds": [os.path.join(current_dir, f"{i}.png") for i in range(1,6)],
            "reaction": os.path.join(current_dir, "reaction.txt"),
            "ranked_preds": os.path.join(current_dir, "ranked_preds.txt"),
            "input": os.path.join(current_dir, "input.txt")
        }

        # 验证文件存在性
        if not all(os.path.exists(f) for f in [required_files["reference"], required_files["reaction_img"]]):
            continue
        
        with open(required_files["reaction"], 'r') as f:
             original_reaction = f.read().strip()
             
        with open(required_files["input"], 'r') as f:
            target_smiles = f.read().strip()
            
        with open(required_files["ranked_preds"], 'r') as f:
            pred_lines = [line.strip() for line in f.readlines()[:5]]

        matchboolindex=0
        # 生成预测反应式示意图
        for idx, precursor in enumerate(pred_lines, 1):
            if not precursor or '>' in precursor:  # 过滤无效行
                continue
                
            full_reaction = f"{precursor}>>{target_smiles}"
            matchbool=are_reactions_equal(original_reaction,full_reaction)
            if matchbool:
                matchboolindex=1
                
                print(f"匹配成功: {precursor}>>{target_smiles}")
                break
            else:
                
                print(f"匹配失败: {precursor}>>{target_smiles}")
                continue
            
        if matchboolindex==1 and matchcount==22:
            
            continue
        if matchboolindex==0 and nomatchcount==164:
            
            continue
        if nomatchcount==164 and matchcount==22:
            break

        # ---------- 开始构建页面 ----------
        # 添加Reference
        with open(required_files["reference"], 'r') as f:
            ref_content = f.read().strip()
            reflist=ref_content.split('Article;')
            for ref in reflist:
                if 'Organic Letters' in ref:
                    ref_content=ref
                    break
        
        index=index+1
        # 主标题
        p = doc.add_paragraph()
        #runner = p.add_run(f"Reaction {index}\n")
        if matchboolindex==1:
            matchcount=matchcount+1
            runner = p.add_run(f"Reaction {index}\nMatch\n")
        else:
            nomatchcount=nomatchcount+1
            runner = p.add_run(f"Reaction {index}\nNot Match\n")
        runner.bold = True
        runner.font.size = Pt(14)

        # Reference部分
        p = doc.add_paragraph()
        p.add_run("Reference: ").bold = True
        p.add_run(ref_content)
        
        # 添加反应式示意图
        doc.add_paragraph().add_run("Reaction from Reference").bold = True
        try:
            doc.add_picture(required_files["reaction_img"], width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"图片加载失败: {str(e)}")

        # 添加预测路径
        doc.add_paragraph().add_run("Top5 of Editretro predict").bold = True
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 插入预测图片（1-5.png）
        valid_pred_count = 0
        for i in range(1, 6):
            img_path = os.path.join(current_dir, f"{i}.png")
            if os.path.exists(img_path):
                try:
                    run = para.add_run()
                    run.add_picture(img_path, width=Inches(5.5))  # 并排显示
                    valid_pred_count += 1
                    # 每两个图片换行
                    if valid_pred_count % 1 == 0:
                        para = doc.add_paragraph()
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    doc.add_paragraph(f"预测路径 {i} 加载失败: {str(e)}")

        # 添加分页符
        doc.add_page_break()

    # 保存文档
    doc.save(output_file)
    print(f"匹配成功数量：{matchcount}")
    print(f"匹配失败数量：{nomatchcount}")
    print(f"报告已生成：{os.path.abspath(output_file)}")

if __name__ == "__main__":
    generate_report(
        source_dir="reactions_data",
        output_file="chemical_report-f.docx"
    )